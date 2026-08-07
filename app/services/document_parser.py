from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

@dataclass
class DocumentSection:
    heading: str
    content: str

def inspect_tables(content_bytes: bytes) -> None:
    document = Document(BytesIO(content_bytes))

    print(f"Table count: {len(document.tables)}")

    for table in document.tables:
        print("TABLE")

        for row in table.rows:
            values = []

            for cell in row.cells:
                values.append(cell.text.strip())

            print(values)

def table_to_text(table: Table) -> str:
    rows: list[str] = []

    # Iterate table.rows
    # Extract every cell's text
    # Join cells using " | "
    # Add each row to rows
    rows.append("[TABLE]")
    for row in table.rows:
        cells: list[str] = []

        for cell in row.cells:
            cells.append(cell.text.strip())

        row_text = " | ".join(cells)
        rows.append(row_text)

    rows.append("[/TABLE]")

    # Finally join rows using "\n"
    return "\n".join(rows)

def iterate_document_blocks(document):
    body = document.element.body

    for child in body.iterchildren():

        if isinstance(child, CT_P):
            # Convert this XML element into Paragraph
            yield Paragraph(child, document)

        elif isinstance(child, CT_Tbl):
            # Convert this XML element into Table
            yield Table(child, document)

def extract_docx_sections(content_bytes: bytes) -> list[DocumentSection]:
    document = Document(BytesIO(content_bytes))

    sections: list[DocumentSection] = []

    current_heading = "Introduction"
    current_content: list[str] = []

    for block in iterate_document_blocks(document):

        if isinstance(block, Paragraph):
            text = block.text.strip()

            if not text:
                continue

            style_name = getattr(block.style, "name", "")

            if style_name.startswith("Heading"):
                # Save previous section if it has content.
                if current_content:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n\n".join(current_content)
                        )
                    )

                # Change current_heading.
                current_heading = text

                # Reset current_content.
                current_content = []

            else:
                # Add paragraph text.
                current_content.append(text)

        elif isinstance(block, Table):
            # Convert table to text.
            # Add the resulting text to current_content.
            table_text = table_to_text(block)
            current_content.append(table_text)

    # Save final section if content exists.
    if current_content:
        sections.append(
            DocumentSection(
                heading=current_heading,
                content="\n\n".join(current_content)
            )
        )

    return sections